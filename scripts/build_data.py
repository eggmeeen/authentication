from __future__ import annotations

import csv
import json
import re
import sqlite3
import sys
import urllib.request
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ModuleNotFoundError as exc:
    raise SystemExit("python-docx is required: python3 -m pip install python-docx") from exc


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT.parent / "认证项目多场所一览表（gpt版）-1.docx"
CITY_GEO_URL = "https://raw.githubusercontent.com/88250/city-geo/master/data.json"
CITY_GEO_PATH = ROOT / "data" / "source" / "city-geo.json"
PUBLIC_DATA_PATH = ROOT / "public" / "data" / "projects.json"
CSV_PATH = ROOT / "data" / "projects.csv"
SQLITE_PATH = ROOT / "data" / "projects.sqlite"

FALLBACK_COORDS = {
    # city-geo leaves this Xinjiang county-level city blank; GeoJSON.CN 1.6.x
    # lists the administrative center as [79.074965, 39.867776].
    ("新疆维吾尔自治区", "自治区直辖县级行政区划", "图木舒克市"): {"lng": 79.074965, "lat": 39.867776},
}


ETHNIC_TOKENS = [
    "蒙古族",
    "蒙古",
    "藏族",
    "藏",
    "回族",
    "回",
    "维吾尔族",
    "维吾尔",
    "哈萨克族",
    "哈萨克",
    "柯尔克孜",
    "朝鲜族",
    "朝鲜",
    "彝族",
    "彝",
    "白族",
    "白",
    "傣族",
    "傣",
    "苗族",
    "苗",
    "侗族",
    "侗",
    "土家族",
    "土家",
    "壮族",
    "壮",
    "布依族",
    "布依",
    "哈尼族",
    "哈尼",
    "傈僳族",
    "傈僳",
    "景颇族",
    "景颇",
    "拉祜族",
    "拉祜",
    "佤族",
    "佤",
    "纳西族",
    "纳西",
    "羌族",
    "羌",
]

DIRECT_ADMIN_CITY_NAMES = {
    "市辖区",
    "省直辖县级行政区划",
    "自治区直辖县级行政区划",
    "县",
}


def download_city_geo() -> None:
    CITY_GEO_PATH.parent.mkdir(parents=True, exist_ok=True)
    if CITY_GEO_PATH.exists() and CITY_GEO_PATH.stat().st_size > 100_000:
        return
    with urllib.request.urlopen(CITY_GEO_URL, timeout=30) as response:
        CITY_GEO_PATH.write_bytes(response.read())


def compact_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()


def clean_address(value: str) -> str:
    text = compact_text(value.replace("\n", " / "))
    text = re.sub(r"^(注册地址|地址|办公地址|办公地点)[:：]\s*", "", text)
    text = re.sub(r"\s*/\s*", " / ", text)
    return text


def clean_company(value: str) -> str:
    text = compact_text(value.replace("\n", " "))
    text = re.sub(r"^\d{1,3}(?=[\u4e00-\u9fffA-Za-z（(])", "", text)
    return re.sub(r"\s*/\s*$", "", text).strip()


def province_short_name(name: str) -> str:
    replacements = [
        ("壮族自治区", ""),
        ("回族自治区", ""),
        ("维吾尔自治区", ""),
        ("自治区", ""),
        ("特别行政区", ""),
        ("省", ""),
        ("市", ""),
    ]
    short = name
    for suffix, replacement in replacements:
        if short.endswith(suffix):
            short = short[: -len(suffix)] + replacement
    return short


def common_city_name(name: str) -> str:
    if name.endswith("自治州"):
        short = name[:-3]
        changed = True
        while changed:
            changed = False
            for token in ETHNIC_TOKENS:
                if short.endswith(token):
                    short = short[: -len(token)]
                    changed = True
        if short:
            return f"{short}州"
    return name


def city_label(name: str) -> str:
    common = common_city_name(name)
    if common.endswith("市"):
        return common[:-1]
    return common


def city_aliases(city: dict[str, str]) -> set[str]:
    name = city["city"]
    province = city["province"]
    aliases: set[str] = set()

    if name == "市辖区":
        aliases.update({province, province_short_name(province)})
        return {alias for alias in aliases if len(alias) >= 2}

    if name in DIRECT_ADMIN_CITY_NAMES:
        return set()

    aliases.add(name)
    common = common_city_name(name)
    aliases.add(common)
    for value in [name, common]:
        for suffix in ["市", "地区", "盟", "州", "自治州"]:
            if value.endswith(suffix):
                base = value[: -len(suffix)]
                aliases.add(base)
                aliases.add(f"{base}{suffix}")
    return {alias for alias in aliases if len(alias) >= 2}


def area_aliases(area_name: str) -> set[str]:
    aliases = {area_name}
    if area_name.endswith("市"):
        aliases.add(f"{area_name[:-1]}县")
        aliases.add(area_name[:-1])
    elif area_name.endswith("县"):
        aliases.add(f"{area_name[:-1]}市")
        aliases.add(area_name[:-1])
    elif area_name.endswith("旗"):
        aliases.add(area_name[:-1])
    elif area_name.endswith("行政委员会"):
        aliases.add(area_name[: -len("行政委员会")])
    return {alias for alias in aliases if len(alias) >= 2}


def as_coord(record: dict[str, str]) -> dict[str, float]:
    if record.get("lng") and record.get("lat"):
        return {"lng": float(record["lng"]), "lat": float(record["lat"])}
    key = (record.get("province", ""), record.get("city", ""), record.get("area", ""))
    if key in FALLBACK_COORDS:
        return FALLBACK_COORDS[key]
    raise ValueError(f"missing coordinates for {key}")


def build_indexes(city_geo: list[dict[str, str]]) -> dict[str, Any]:
    provinces: dict[str, dict[str, str]] = {}
    province_aliases: list[tuple[str, str]] = []
    city_records: list[dict[str, str]] = []
    area_records: list[dict[str, str]] = []

    for record in city_geo:
        province = record["province"]
        provinces.setdefault(province, record)
        if record["area"] == "":
            city_records.append(record)
        elif record["area"] not in {"市辖区", "县"}:
            area_records.append(record)

    for province in provinces:
        aliases = {province, province_short_name(province)}
        if province == "内蒙古自治区":
            aliases.add("内蒙古")
        elif province == "广西壮族自治区":
            aliases.add("广西")
        elif province == "宁夏回族自治区":
            aliases.add("宁夏")
        elif province == "新疆维吾尔自治区":
            aliases.add("新疆")
        elif province == "西藏自治区":
            aliases.add("西藏")
        province_aliases.extend((alias, province) for alias in aliases if alias)

    city_alias_index: list[tuple[str, dict[str, str]]] = []
    for city in city_records:
        city_alias_index.extend((alias, city) for alias in city_aliases(city))

    area_alias_index: list[tuple[str, dict[str, str]]] = []
    for area in area_records:
        area_alias_index.extend((alias, area) for alias in area_aliases(area["area"]))

    return {
        "province_aliases": sorted(province_aliases, key=lambda item: len(item[0]), reverse=True),
        "city_aliases": sorted(city_alias_index, key=lambda item: len(item[0]), reverse=True),
        "area_aliases": sorted(area_alias_index, key=lambda item: len(item[0]), reverse=True),
        "city_center": {(record["province"], record["city"]): record for record in city_records},
    }


def infer_province(text: str, indexes: dict[str, Any], prefix_only: bool = False) -> str:
    for alias, province in indexes["province_aliases"]:
        if prefix_only and text.startswith(alias):
            return province
        if not prefix_only and alias in text:
            return province
    return ""


def find_city(text: str, province: str, indexes: dict[str, Any]) -> tuple[dict[str, str] | None, str]:
    for alias, city in indexes["city_aliases"]:
        if province and city["province"] != province:
            continue
        if alias in text:
            return city, alias
    return None, ""


def find_area(
    text: str,
    province: str,
    city_name: str,
    indexes: dict[str, Any],
) -> tuple[dict[str, str] | None, str]:
    for alias, area in indexes["area_aliases"]:
        if province and area["province"] != province:
            continue
        if city_name and area["city"] != city_name:
            continue
        if alias in text:
            return area, alias
    return None, ""


def display_city_name(province: str, city_name: str, area_name: str) -> str:
    if city_name == "市辖区":
        return province
    if city_name in {"省直辖县级行政区划", "自治区直辖县级行政区划", "县"} and area_name:
        return area_name
    return city_name


def match_location(row: dict[str, str], indexes: dict[str, Any]) -> dict[str, Any]:
    address = row["address"]
    region = row["region"]
    company = row["company"]
    text_all = f"{address} {region} {company}"

    province = (
        infer_province(address, indexes, prefix_only=True)
        or infer_province(region, indexes)
        or infer_province(address, indexes)
        or infer_province(text_all, indexes)
    )
    city_record, city_alias = find_city(address, province, indexes)
    if not city_record:
        city_record, city_alias = find_city(f"{address} {region}", province, indexes)
    if not city_record:
        city_record, city_alias = find_city(text_all, province, indexes)

    city_name = city_record["city"] if city_record else ""
    area_record, area_alias = find_area(address, province, city_name, indexes)
    if not area_record:
        area_record, area_alias = find_area(text_all, province, city_name, indexes)

    if not city_record and area_record:
        province = province or area_record["province"]
        city_name = area_record["city"]
        city_record = indexes["city_center"].get((area_record["province"], area_record["city"]))
    elif city_record:
        province = province or city_record["province"]

    display_city = display_city_name(province, city_name, area_record["area"] if area_record else "")
    if not city_record and area_record:
        city_coord_record = area_record
    elif city_name in {"省直辖县级行政区划", "自治区直辖县级行政区划", "县"} and area_record:
        city_coord_record = area_record
    else:
        city_coord_record = city_record

    exact_coord_record = area_record or city_coord_record
    if not display_city or not city_coord_record or not exact_coord_record:
        return {
            "province": province,
            "prefecture": city_name,
            "displayCity": display_city,
            "displayCityLabel": city_label(display_city) if display_city else "",
            "district": area_record["area"] if area_record else "",
            "cityCoord": None,
            "addressCoord": None,
            "coordinateLevel": "unresolved",
            "matchMethod": "unresolved",
            "matchedAlias": city_alias or area_alias,
        }

    if area_record:
        method = "district"
        level = "district"
    else:
        method = "prefecture"
        level = "prefecture"

    return {
        "province": province,
        "prefecture": city_name,
        "prefectureCommon": common_city_name(display_city),
        "displayCity": display_city,
        "displayCityLabel": city_label(display_city),
        "district": area_record["area"] if area_record else "",
        "cityCoord": as_coord(city_coord_record),
        "addressCoord": as_coord(exact_coord_record),
        "coordinateLevel": level,
        "matchMethod": method,
        "matchedAlias": area_alias or city_alias,
    }


def extract_rows() -> tuple[str, list[dict[str, str]]]:
    if not SOURCE_DOCX.exists():
        raise SystemExit(f"Source DOCX not found: {SOURCE_DOCX}")
    document = Document(SOURCE_DOCX)
    source_title = next((compact_text(p.text) for p in document.paragraphs if compact_text(p.text)), "")
    if not document.tables:
        raise SystemExit("Source DOCX contains no tables")
    table = document.tables[0]
    rows: list[dict[str, str]] = []
    for index, table_row in enumerate(table.rows[1:], start=1):
        raw_cells = [cell.text.strip() for cell in table_row.cells]
        cells = [
            compact_text(raw_cells[0].replace("\n", " / ")),
            compact_text(raw_cells[1].replace("\n", " / ")),
            compact_text(raw_cells[2].replace("\n", " ")),
            compact_text(raw_cells[3].replace("\n", " / ")),
        ]
        if len(cells) < 4:
            continue
        company_raw = cells[2]
        address_raw = cells[3]
        if not company_raw and not address_raw:
            continue
        rows.append(
            {
                "id": str(index),
                "region": cells[0],
                "sequence": cells[1],
                "companyRaw": company_raw,
                "company": clean_company(company_raw),
                "addressRaw": address_raw,
                "address": clean_address(address_raw),
            }
        )
    return source_title, rows


def build_dataset() -> dict[str, Any]:
    download_city_geo()
    city_geo = json.loads(CITY_GEO_PATH.read_text(encoding="utf-8"))
    indexes = build_indexes(city_geo)
    source_title, raw_rows = extract_rows()

    projects: list[dict[str, Any]] = []
    unresolved: list[int] = []
    for row in raw_rows:
        location = match_location(row, indexes)
        project = {
            "id": int(row["id"]),
            "region": row["region"],
            "sequence": row["sequence"],
            "companyRaw": row["companyRaw"],
            "company": row["company"],
            "addressRaw": row["addressRaw"],
            "address": row["address"],
            **location,
        }
        if project["coordinateLevel"] == "unresolved":
            unresolved.append(project["id"])
        projects.append(project)

    city_groups: dict[str, dict[str, Any]] = {}
    for project in projects:
        if not project["cityCoord"]:
            continue
        key = f"{project['province']}|{project['displayCity']}"
        group = city_groups.setdefault(
            key,
            {
                "key": key,
                "province": project["province"],
                "city": project["displayCity"],
                "label": project["displayCityLabel"],
                "lng": project["cityCoord"]["lng"],
                "lat": project["cityCoord"]["lat"],
                "count": 0,
                "projectIds": [],
            },
        )
        group["count"] += 1
        group["projectIds"].append(project["id"])

    methods = Counter(project["matchMethod"] for project in projects)
    provinces = Counter(project["province"] for project in projects if project["province"])
    dataset = {
        "metadata": {
            "title": "认证项目多场所地图",
            "sourceFile": str(SOURCE_DOCX),
            "sourceTitle": source_title,
            "extractedProjects": len(projects),
            "uniquePlaces": len(city_groups),
            "generatedAt": datetime.now(timezone.utc).isoformat(),
            "coordinateSource": "city-geo data.json; GeoJSON.CN fallback for 图木舒克市",
            "coordinateSourceUrl": CITY_GEO_URL,
            "fallbackCoordinateSourceUrl": "https://geojson.cn/data/atlas/china",
            "coordinateNote": "当前坐标为地级行政区或区县行政中心点，用于地图定位和跳转；不是门牌级精确坐标。",
            "unresolvedIds": unresolved,
            "matchMethods": dict(methods),
            "provinceCounts": dict(provinces),
        },
        "projects": projects,
        "cities": sorted(city_groups.values(), key=lambda item: (-item["count"], item["province"], item["city"])),
    }
    return dataset


def write_csv(projects: list[dict[str, Any]]) -> None:
    CSV_PATH.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "id",
        "region",
        "company",
        "address",
        "province",
        "displayCity",
        "displayCityLabel",
        "district",
        "cityLng",
        "cityLat",
        "addressLng",
        "addressLat",
        "coordinateLevel",
        "matchMethod",
        "matchedAlias",
    ]
    with CSV_PATH.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        for project in projects:
            writer.writerow(
                {
                    "id": project["id"],
                    "region": project["region"],
                    "company": project["company"],
                    "address": project["address"],
                    "province": project["province"],
                    "displayCity": project["displayCity"],
                    "displayCityLabel": project["displayCityLabel"],
                    "district": project["district"],
                    "cityLng": project["cityCoord"]["lng"] if project["cityCoord"] else "",
                    "cityLat": project["cityCoord"]["lat"] if project["cityCoord"] else "",
                    "addressLng": project["addressCoord"]["lng"] if project["addressCoord"] else "",
                    "addressLat": project["addressCoord"]["lat"] if project["addressCoord"] else "",
                    "coordinateLevel": project["coordinateLevel"],
                    "matchMethod": project["matchMethod"],
                    "matchedAlias": project["matchedAlias"],
                }
            )


def write_sqlite(dataset: dict[str, Any]) -> None:
    SQLITE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if SQLITE_PATH.exists():
        SQLITE_PATH.unlink()
    conn = sqlite3.connect(SQLITE_PATH)
    conn.execute(
        """
        CREATE TABLE projects (
            id INTEGER PRIMARY KEY,
            region TEXT,
            company_raw TEXT,
            company TEXT NOT NULL,
            address_raw TEXT,
            address TEXT NOT NULL,
            province TEXT,
            prefecture TEXT,
            display_city TEXT,
            display_city_label TEXT,
            district TEXT,
            city_lng REAL,
            city_lat REAL,
            address_lng REAL,
            address_lat REAL,
            coordinate_level TEXT,
            match_method TEXT,
            matched_alias TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE cities (
            key TEXT PRIMARY KEY,
            province TEXT,
            city TEXT,
            label TEXT,
            lng REAL,
            lat REAL,
            project_count INTEGER
        )
        """
    )
    conn.execute(
        """
        CREATE VIRTUAL TABLE project_search USING fts5(
            company,
            address,
            province,
            display_city,
            district,
            content='projects',
            content_rowid='id'
        )
        """
    )

    for project in dataset["projects"]:
        conn.execute(
            """
            INSERT INTO projects VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                project["id"],
                project["region"],
                project["companyRaw"],
                project["company"],
                project["addressRaw"],
                project["address"],
                project["province"],
                project["prefecture"],
                project["displayCity"],
                project["displayCityLabel"],
                project["district"],
                project["cityCoord"]["lng"] if project["cityCoord"] else None,
                project["cityCoord"]["lat"] if project["cityCoord"] else None,
                project["addressCoord"]["lng"] if project["addressCoord"] else None,
                project["addressCoord"]["lat"] if project["addressCoord"] else None,
                project["coordinateLevel"],
                project["matchMethod"],
                project["matchedAlias"],
            ),
        )
        conn.execute(
            """
            INSERT INTO project_search(rowid, company, address, province, display_city, district)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                project["id"],
                project["company"],
                project["address"],
                project["province"],
                project["displayCity"],
                project["district"],
            ),
        )

    for city in dataset["cities"]:
        conn.execute(
            "INSERT INTO cities VALUES (?, ?, ?, ?, ?, ?, ?)",
            (city["key"], city["province"], city["city"], city["label"], city["lng"], city["lat"], city["count"]),
        )
    conn.execute("CREATE INDEX idx_projects_company ON projects(company)")
    conn.execute("CREATE INDEX idx_projects_city ON projects(display_city)")
    conn.commit()
    conn.close()


def main() -> None:
    dataset = build_dataset()
    PUBLIC_DATA_PATH.parent.mkdir(parents=True, exist_ok=True)
    PUBLIC_DATA_PATH.write_text(json.dumps(dataset, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(dataset["projects"])
    write_sqlite(dataset)

    metadata = dataset["metadata"]
    print(f"wrote {PUBLIC_DATA_PATH}")
    print(f"wrote {CSV_PATH}")
    print(f"wrote {SQLITE_PATH}")
    print(f"projects: {metadata['extractedProjects']}")
    print(f"unique places: {metadata['uniquePlaces']}")
    print(f"match methods: {metadata['matchMethods']}")
    if metadata["unresolvedIds"]:
        print(f"unresolved ids: {metadata['unresolvedIds']}", file=sys.stderr)
        raise SystemExit(2)


if __name__ == "__main__":
    main()
